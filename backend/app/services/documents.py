"""Safe, local plain-text extraction for supported document sources."""

from __future__ import annotations

import zipfile
from pathlib import Path

from app.config import get_settings

MAX_ZIP_ENTRIES = 2_000
MAX_ZIP_ENTRY_BYTES = 25 * 1024 * 1024
MAX_ZIP_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_COMPRESSION_RATIO = 100


class DocumentExtractionError(ValueError):
    """A safe error suitable for mapping to a calm API message."""


def _meaningful(text: str) -> str:
    if not text.strip():
        raise DocumentExtractionError("No extractable text was found. OCR is not supported.")
    if len(text) > get_settings().max_text_chars:
        raise DocumentExtractionError("Extracted text exceeds the configured limit.")
    return text


def _text(path: Path) -> str:
    try:
        raw = path.read_bytes()
        if b"\x00" in raw:
            raise UnicodeDecodeError("utf-8", raw, 0, 1, "NUL byte")
        return _meaningful(raw.decode("utf-8-sig"))
    except (OSError, UnicodeDecodeError) as exc:
        raise DocumentExtractionError("This text document is not valid UTF-8 text.") from exc


def _preflight_docx(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_ZIP_ENTRIES:
                raise DocumentExtractionError("Document archive has too many entries.")
            total = 0
            for info in infos:
                total += info.file_size
                if info.file_size > MAX_ZIP_ENTRY_BYTES or total > MAX_ZIP_UNCOMPRESSED_BYTES:
                    raise DocumentExtractionError(
                        "Document archive is too large to extract safely."
                    )
                if (
                    info.compress_size
                    and info.file_size / info.compress_size > MAX_COMPRESSION_RATIO
                ):
                    raise DocumentExtractionError(
                        "Document archive has a suspicious compression ratio."
                    )
            if (
                "[Content_Types].xml" not in archive.namelist()
                or "word/document.xml" not in archive.namelist()
            ):
                raise DocumentExtractionError("This is not a valid DOCX document.")
    except zipfile.BadZipFile as exc:
        raise DocumentExtractionError("This is not a valid DOCX document.") from exc


def _docx(path: Path) -> str:
    _preflight_docx(path)
    try:
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            parts.extend("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return _meaningful("\n\n".join(parts))
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError("This DOCX document could not be read safely.") from exc


def _pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        if reader.is_encrypted:
            raise DocumentExtractionError("Password-protected PDFs are not supported.")
        pages = [page.extract_text() or "" for page in reader.pages]
        return _meaningful("\n\n--- Page ---\n\n".join(pages))
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError("This PDF could not be read safely.") from exc


def extract_text(path: Path, extension: str) -> str:
    """Extract local plain text only; callers run this in the durable worker."""
    if extension in {".txt", ".md"}:
        return _text(path)
    if extension == ".pdf":
        return _pdf(path)
    if extension == ".docx":
        return _docx(path)
    raise DocumentExtractionError("Unsupported document format.")
